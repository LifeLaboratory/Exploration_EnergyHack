import docx
from base.base_sql import Sql
from word.word_sql import *


def find_vl_kl(find_kab):
    kl = 0
    vl = 0
    for kab in find_kab:
        if "каб" in kab["Название"].lower() and "каб.оп" not in kab["Название"].lower() and "каб. оп" not in kab[
            "Название"].lower():
            try:
                try:
                    kl += float(kab["Длина_по_цепям"])
                except:
                    kl += float(kab["Длина_участка_по_цепям"])
            except:
                pass
        else:
            try:
                try:
                    vl += float(kab["Длина_по_цепям"])
                except:
                    vl += float(kab["Длина_участка_по_цепям"])
            except:
                pass
    return kl, vl


def create_word(comp_id):
    sql_answer = Sql.exec(SQL_NAME_COMPANY.format(id=comp_id))
    try:
        name_company = sql_answer[0]["Название"]
    except:
        return "Данной компании не существует"
    doc = docx.Document()
    header = doc.add_paragraph("")
    header.add_run("Анализ технического состояния и возрастная структура линий электропередачи и подстанций").bold = True
    year = 2020
    date = "01.01.{}".format(year)
    sql_answer = Sql.exec(SQL_FIND_KV.format(id=comp_id))
    try:
        kv = sql_answer[0]["Напряжение"]
    except:
        kv = sql_answer[0]["Напряжение"]
    first_str = doc.add_paragraph("")
    first_str.add_run(name_company).bold = True
    str1 = "Протяженность ВЛ {kv} кВ и КЛ {kv} кВ, количество и суммарная мощность ПС {kv} кВ, находящихся в собственности {name_company}, по состоянию на {date} г. составили:"
    doc.add_paragraph(str1.format(kv=kv, date=date, name_company=name_company))
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    find_kab = Sql.exec(SQL_FIND_KAB_DL.format(id=comp_id))
    kl, vl = find_vl_kl(find_kab)
    cell = table.cell(0, 0)
    cell.text = "Протяженность действующих ВЛ и КЛ\n(в одноцепном исчислении), км"
    cell = table.cell(1, 0)
    cell.text = "Количество и суммарная установленная\nмощность ПС, шт./МВА"
    cell = table.cell(0, 1)
    cell.text = "ВЛ – {vl}\nКЛ – {kl}".format(vl=int(vl*100)/100, kl=int(kl*100)/100)
    mva = 0
    sql_answers = Sql.exec(SQL_FIND_ALL_MOSH.format(id=comp_id))
    for mosh in sql_answers:
        try:
            mva += float(mosh["Номинальная_мощность"])
        except:
            pass
    sht = len(Sql.exec(SQL_FIND_ALL_PODST.format(id=comp_id)))
    cell = table.cell(1, 1)
    cell.text = "{sht} / {mva}".format(sht=sht, mva=int(mva*100)/100)
    doc.add_paragraph("")
    str = "Далее приведена возрастная структура линий электропередачи и подстанций {kv} кВ {name_company} по состоянию на {date} г. с разбивкой по электросетевым предприятиям. "
    doc.add_paragraph(str1.format(kv=kv, date=date, name_company=name_company))
    sql_answers = Sql.exec(SQL_FIND_SETT.format(id=comp_id))
    for sql_ans in sql_answers:
        name_predpr = sql_ans["Название"]
        first_str = doc.add_paragraph("")
        first_str.add_run("Филиал {nc} {np}".format(nc=name_company, np=name_predpr)).bold = True
        str = "Протяженность ВЛ {kv} кВ и КЛ {kv} кВ, количество и суммарная мощность ПС {kv} кВ,  обслуживаемых {np} по состоянию на {date} г. составили:"
        doc.add_paragraph(str.format(kv=kv, np=name_predpr, date=date))
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell.text = "Протяженность действующих ВЛ и КЛ\n(в одноцепном исчислении), км"
        cell = table.cell(1, 0)
        cell.text = "Количество и суммарная установленная\nмощность ПС, шт./МВА"
        find_kab = Sql.exec(SQL_FIND_KAB_SET_DL.format(id=sql_ans["id"]))
        kl, vl = find_vl_kl(find_kab)
        cell = table.cell(0, 1)
        cell.text = "ВЛ – {vl}\nКЛ – {kl}".format(vl=int(vl*100)/100, kl=int(kl*100)/100)
        mva = 0
        sql_answers = Sql.exec(SQL_FIND_SETT_MOSH.format(id=sql_ans["id"]))
        for mosh in sql_answers:
            try:
                mva += float(mosh["Номинальная_мощность"])
            except:
                pass
        sht = len(Sql.exec(SQL_FIND_SETT_PODST.format(id=sql_ans["id"])))
        cell = table.cell(1, 1)
        cell.text = "{sht} / {mva}".format(sht=sht, mva=int(mva * 100) / 100)
        text = "Анализ технического состояния электросетевых объектов напряжением {kv} кВ {np} показал:"
        doc.add_paragraph(text.format(kv=kv, np=name_predpr))
        text = "•	{kolvo_podst} подстанций ({proc}% от общего числа ПС {kv} кВ) отработали более 50 лет;"
        kolvo_podst = len(Sql.exec(SQL_FIND_PODST_OLD.format(id=sql_ans["id"], year=year)))
        proc = 0
        if kolvo_podst > 0:
            proc = int(kolvo_podst/sht * 10000)/100
        doc.add_paragraph(text.format(kolvo_podst=kolvo_podst, proc=proc, kv=kv))
        text = "•	{kolvo_mosh} МВА трансформаторной мощности ({proc}% от общей трансформаторной мощности напряжением {kv} кВ) отработало более 50 лет;"
        kolvo_mosh = len(Sql.exec(SQL_FIND_PODST_OLD.format(id=sql_ans["id"], year=year)))
        proc = 0
        if kolvo_mosh > 0:
            proc = int(kolvo_mosh / mva * 10000) / 100
        doc.add_paragraph(text.format(kolvo_mosh=kolvo_mosh, proc=proc, kv=kv))
        text = "•	воздушные линии электропередачи {kv} кВ протяженностью {voz_dl} км в одноцепном исчислении ({proc}% от общей протяженности ВЛ {kv} кВ) отработали более 50 лет;"
        text1 = "•	кабельные линии электропередачи {kv} кВ протяженностью {kab_dl} км ({proc}% от общей протяженности КЛ {kv} кВ) находятся в эксплуатации до 35 лет."
        find_kab = Sql.exec(SQL_FIND_KAB_SET_OLD.format(id=sql_ans["id"], year=year))
        kls, vls = find_vl_kl(find_kab)
        proc = 0
        if vls > 0:
            proc = int(vls / vl * 10000) / 100
        doc.add_paragraph(text.format(voz_dl=int(vls*100)/100, proc=proc, kv=kv))
        proc = 0
        if kls > 0:
            proc = int(kls / kl * 10000) / 100
        doc.add_paragraph(text1.format(kab_dl=kls, proc=proc, kv=kv))
    doc.save('static/Итоговый_отчёт.docx')
